# -*- coding: utf-8 -*-
"""智慧社区物业管理系统 - 课程综合报告生成脚本"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_PATH = r'E:\软件综合课程设计\课程综合报告.docx'

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

for lvl, sz, sp in [('Heading 1', 18, 18), ('Heading 2', 15, 14), ('Heading 3', 13, 12)]:
    s = doc.styles[lvl]
    s.font.name = '黑体'
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s.font.size = Pt(sz)
    s.paragraph_format.space_before = Pt(sp)
    s.paragraph_format.space_after = Pt(sp)

def sr(run, fn='宋体', sz=Pt(12), b=False, c=None):
    run.font.name = fn; run.font.size = sz; run.bold = b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if c: run.font.color.rgb = c

def body(text, indent=True):
    p = doc.add_paragraph()
    r = p.add_run(text); sr(r)
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    return p

def bodyb(text):
    p = doc.add_paragraph()
    r = p.add_run(text); sr(r, b=True)
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h); sr(r, '宋体', Pt(10), True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 灰色表头背景
        from docx.oxml import OxmlElement
        tcPr = c._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        tcPr.append(shading)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(v)); sr(r, '宋体', Pt(10))
    return t

def bullet(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(it); sr(r)

def numbered(items):
    for it in items:
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(it); sr(r)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

# ══════════════════════════════════════════════════
#  封面
# ══════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('应用软件综合课程设计 I'); sr(r, '黑体', Pt(26), True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('实验报告——课程综合报告'); sr(r, '黑体', Pt(20), True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('智慧社区物业管理系统'); sr(r, '黑体', Pt(22), True, RGBColor(0x1A, 0x56, 0xDB))

doc.add_paragraph()

info = [
    ('专业年级：', '软件工程专业2023级'),
    ('组    长：', '蒙焕好（2023112596）'),
    ('组    员：', '农麒民（2023112563）、刘向阳（2023112568）'),
    ('指导老师（职称）：', '韩敏（教授）'),
    ('提交日期：', '2026年6月'),
]
for label, val in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label); sr(r1, '宋体', Pt(14), True)
    r2 = p.add_run(val); sr(r2, '宋体', Pt(14))

doc.add_page_break()

# ══════════════════════════════════════════════════
#  第一章 系统概述
# ══════════════════════════════════════════════════
H('一、系统概述', 1)

H('1.1 项目背景', 2)
body('随着我国城市化进程的加速推进，住宅小区规模不断扩大，传统物业管理模式面临效率低下、信息不透明、业主体验差等突出问题。物业管理涉及房屋楼栋管理、车位资源分配、费用收缴、设施维护、门禁安全等多个维度，依赖人工台账和纸质单据的方式已难以满足现代化社区的管理需求。')
body('在此背景下，本课题组设计并开发了"智慧社区物业管理系统"，旨在通过信息化手段实现物业管理的数字化、规范化和可视化，提升物业运营效率，改善业主服务体验。')

H('1.2 系统目标', 2)
body('本系统的核心建设目标包括：')
numbered([
    '构建统一的物业管理平台，覆盖房屋楼栋、车位、费用、设施、门禁、反馈等核心业务模块；',
    '实现业主端自助服务，支持在线缴费、设施报修、设施借用、车位选购、留言反馈等操作；',
    '提供管理端高效运营工具，包括数据统计概览、账单批量生成、门禁卡全生命周期管理、进出记录追踪等；',
    '引入 3D 可视化技术，实现楼栋三维模型展示和车位可视化选购；',
    '基于 JWT 无状态认证机制保障系统安全性，实现角色级别的权限控制。',
])

H('1.3 系统问题定义', 2)
body('在项目启动阶段，我们通过调研明确了以下核心问题：')
bullet([
    '信息孤岛问题：房屋、车位、业主等信息分散在不同台账中，缺乏统一数据管理平台；',
    '缴费效率低：物业费、车位管理费等依赖人工催收和线下支付，对账困难；',
    '车位管理混乱：车位状态不透明，业主无法直观了解车位分布和剩余情况；',
    '设施管理缺失：公共设施借用缺乏线上化流程，设施维护记录无法追溯；',
    '门禁安全隐患：传统门禁卡管理依赖人工登记，无法追踪进出记录；',
    '沟通渠道不畅：业主投诉、建议缺乏规范化反馈通道，物业无法及时响应。',
])

H('1.4 系统范围', 2)
body('本系统覆盖智慧社区物业管理的全业务链条，按角色划分为两大子系统：')
bodyb('业主端（OWNER）：')
body('提供首页概览、在线缴费、车位选购、设施借用、设施报修、留言反馈、门禁卡查看、社区公告、个人中心等9个功能页面，满足业主日常自助服务需求。')
bodyb('管理端（ADMIN）：')
body('提供数据概览（Dashboard）、用户管理、费用项目管理、账单管理、楼栋房屋管理、楼栋平面图、车位管理、设施管理、报修管理、留言反馈管理、门禁卡管理、进出记录管理、公告管理等13个管理页面，实现物业运营的全流程管控。')

doc.add_page_break()

# ══════════════════════════════════════════════════
#  第二章 系统需求分析
# ══════════════════════════════════════════════════
H('二、系统需求分析', 1)

H('2.1 系统流程概述', 2)
body('本系统采用 B/S 架构，用户通过浏览器访问系统。整体业务流程如下：')
body('① 用户注册/登录 → ② 角色识别（业主/管理员）→ ③ 进入对应工作台 → ④ 执行各业务操作 → ⑤ 数据持久化 → ⑥ 结果反馈')
body('')
body('系统核心业务流程图（文本形式）：')
code('''
┌─────────┐    ┌──────────┐    ┌─────────────┐
│  注册    │───→│  登录    │───→│  JWT认证    │
│ Register │    │  Login   │    │  JwtFilter  │
└─────────┘    └──────────┘    └──────┬──────┘
                                      │
                    ┌─────────────────┼──────────────────┐
                    ▼                                    ▼
          ┌──────────────────┐               ┌──────────────────┐
          │   业主端工作台    │               │   管理端工作台    │
          │   Owner Layout   │               │   Admin Layout   │
          └────────┬─────────┘               └────────┬─────────┘
                   │                                   │
    ┌──────┬───────┼──────┬──────┬──────┐   ┌────┬────┼────┬────┬─────┐
    ▼      ▼       ▼      ▼      ▼      ▼   ▼    ▼    ▼    ▼    ▼     ▼
  缴费  车位    设施   报修  反馈  门禁  概览  用户  费用  楼栋  车位  ...
''')

H('2.2 业主端业务分析', 2)

H('2.2.1 在线缴费', 3)
body('业主登录后进入在线缴费页面，系统自动展示其名下所有房屋和车位关联的待缴账单。支持按状态筛选（待缴费/已缴费/已逾期），点击"立即缴费"完成模拟支付。支付成功后账单状态实时更新。')
body('核心业务规则：')
bullet([
    '账单编号格式：yyyyMMdd-P + 业主ID三位补零（如 20260601-P003）',
    '仅 PENDING 状态可支付，已支付账单不可重复操作',
    '逾期账单（due_date < 当前日期 且 status=PENDING）在统计中单独标记',
])

H('2.2.2 车位选购', 3)
body('业主进入车位管理页面，系统以 3D 可视化方式展示各楼栋车位分布。业主可点击楼栋切换视图，查看空闲车位（绿色）、已售车位（灰色）、已锁定车位（红色）。点击空闲车位查看详细信息，确认后完成购买。')
body('核心业务规则：')
bullet([
    '车位状态：FREE（空闲）→ SOLD（已售）/ LOCKED（锁定）/ RESERVED（预订）',
    '仅 FREE 状态车位可购买，已售车位不可重复购买',
    '购买操作在事务中执行，同时更新车位状态和购买信息',
    '业主可释放已购车位，释放后状态恢复为 FREE',
])

H('2.2.3 设施借用', 3)
body('业主浏览社区公共设施列表（运动器材、工具设备等），按分类筛选并提交借用申请。管理员审批通过后，设施状态自动锁定为 BOOKED。业主使用完毕后在线归还，设施恢复可借用状态。')
body('状态流转：AVAILABLE → BOOKED（审批通过）→ AVAILABLE（归还）')

H('2.2.4 设施报修', 3)
body('业主在线提交设施报修申请，填写故障类型（水电/门窗/电梯/公共设施/其他）、描述和故障图片。报修单号由系统自动生成，格式为 yyyyMMdd-R + 业主ID三位补零。管理员可查看并处理报修工单。')

H('2.2.5 留言反馈', 3)
body('业主通过留言反馈模块向物业提交建议、投诉或咨询。支持反馈类型筛选、状态筛选。管理员查看后可回复，业主也可追加回复，形成双向沟通时间线。')
body('状态流转：PENDING → PROCESSING（业主追加回复）→ REPLIED（管理员回复）→ CLOSED（管理员关闭）')

H('2.2.6 个人中心', 3)
body('业主可查看和编辑个人信息（真实姓名、手机号、邮箱、性别、生日、身份证号等），修改密码，上传头像，查看名下房屋列表。')

H('2.3 管理端业务分析', 2)

H('2.3.1 数据概览（Dashboard）', 3)
body('管理员登录后默认进入 Dashboard 页面，顶部展示四项核心指标卡片（用户总数、待办报修数、房屋总数、空闲车位数），中部通过 ECharts 渲染缴费状态柱状图和房屋入住率饼图，底部提供快捷操作入口。')

H('2.3.2 用户管理', 3)
body('管理员对用户进行全生命周期管理，包括新增（自动生成默认密码）、编辑、删除、启用/禁用。支持四表联查（sys_user → community_house → community_unit → community_building），一次查询获取用户及其关联的房屋、楼栋信息。')

H('2.3.3 费用与账单管理', 3)
body('管理员维护物业费用项目（名称、标准金额、收费周期），支持按费用项目批量生成账单。系统自动识别收费维度（房屋类 vs 车位类），加载对应可收费对象，并内置防重复机制避免同一周期重复生成。')

H('2.3.4 楼盘管理', 3)
body('采用楼栋→单元→房屋三级层级结构管理。管理员通过级联选择操作楼栋、单元和房屋，支持办理入住（绑定业主）和退租（解绑业主）。楼栋平面图以可视化方式展示各楼层房屋的入住状态。')

H('2.3.5 门禁卡管理', 3)
body('管理员对门禁卡实施全生命周期管理：发行（自动编号 AC + 日期 + 随机数）→ 编辑权限 → 挂失 → 恢复 → 注销 → 删除。支持卡片类型（业主卡/家庭卡/访客卡/临时卡）、通行楼栋、有效期等多维度配置。')

H('2.3.6 进出记录管理', 3)
body('系统自动记录每次门禁通行事件（卡号、人员、方向、闸机位置、通行状态、拒绝原因），管理员可按多维度条件查询日志，并通过统计面板查看今日通行量、拒绝次数、近7天趋势和各点位统计。')

doc.add_page_break()

# ══════════════════════════════════════════════════
#  第三章 数据库设计
# ══════════════════════════════════════════════════
H('三、数据库设计', 1)

H('3.1 数据库选型', 2)
body('本系统采用 Microsoft SQL Server 作为关系型数据库管理系统，主要原因如下：')
bullet([
    'SQL Server 提供了完善的企业级数据管理能力和高可用性支持；',
    '与 Java 生态通过 mssql-jdbc 驱动无缝集成，支持事务管理和连接池；',
    '丰富的数据类型支持（如 DATETIME、DECIMAL、VARCHAR 等）满足物业管理业务需求。',
])

H('3.2 数据库总体结构', 2)
body('系统数据库名为 community_db，包含以下 15 张核心业务表：')

tbl(['序号','表名','中文名称','主要用途'], [
    ('1', 'sys_user', '用户表', '存储用户账号、密码（BCrypt）、角色、个人信息'),
    ('2', 'community_building', '楼栋表', '楼栋编号、名称、总楼层'),
    ('3', 'community_unit', '单元表', '单元编号，关联楼栋'),
    ('4', 'community_house', '房屋表', '房间号、面积、状态、业主绑定'),
    ('5', 'parking_space', '车位表', '车位编号、楼栋、区域、行列、状态、价格'),
    ('6', 'property_fee_item', '费用项目表', '费用名称、标准金额、收费周期'),
    ('7', 'payment_bill', '缴费账单表', '业主、房屋/车位、费用项、金额、状态、到期日'),
    ('8', 'facility', '设施表', '设施名称、分类、图片、位置、押金、状态'),
    ('9', 'facility_booking', '借用申请表', '设施、业主、用途、时长、审批状态'),
    ('10', 'repair_request', '报修工单表', '业主、类型、描述、图片、处理状态'),
    ('11', 'feedback', '留言反馈表', '业主、类型、标题、内容、状态'),
    ('12', 'feedback_reply', '反馈回复表', '反馈ID、回复者ID、角色、内容'),
    ('13', 'access_card', '门禁卡表', '卡号、业主、类型、状态、通行楼栋、有效期'),
    ('14', 'access_log', '进出记录表', '卡号、人员、方向、闸机、楼栋、通行状态'),
    ('15', 'announcement', '公告表', '标题、内容、类型、发布者、发布状态'),
])

H('3.3 实体关系模型（ER 模型）', 2)
body('系统核心实体关系如下：')
code('''
┌──────────────┐     1:N     ┌──────────────┐
│ sys_user     │◄────────────│community_house│
│ (业主)       │  owner_id   │ (房屋)       │
└──────┬───────┘             └──────┬───────┘
       │                            │
       │ 1:N                        │ N:1
       ▼                            ▼
┌──────────────┐             ┌──────────────┐
│payment_bill  │  N:1        │community_unit│
│ (缴费账单)   │─────────────│ (单元)       │
└──────────────┘             └──────┬───────┘
                                    │ N:1
                                    ▼
                             ┌──────────────┐
                             │community_bldg │
                             │ (楼栋)        │
                             └──────┬───────┘
                                    │ 1:N
                                    ▼
                             ┌──────────────┐
                             │parking_space  │
                             │ (车位)        │
                             └──────────────┘

┌──────────────┐  1:N  ┌──────────────┐
│ sys_user     │───────│ access_card  │
│ (业主)       │       │ (门禁卡)     │
└──────┬───────┘       └──────┬───────┘
       │ 1:N                  │ 1:N
       ▼                      ▼
┌──────────────┐       ┌──────────────┐
│ feedback     │       │ access_log   │
│ (反馈)       │       │ (进出记录)   │
└──────┬───────┘       └──────────────┘
       │ 1:N
       ▼
┌──────────────┐
│feedback_reply│
│ (反馈回复)   │
└──────────────┘

┌──────────────┐  1:N  ┌──────────────┐
│ facility     │───────│facility_booking│
│ (设施)       │       │ (借用申请)    │
└──────────────┘       └──────────────┘
''')

H('3.4 核心表结构设计', 2)

H('3.4.1 sys_user（用户表）', 3)
tbl(['字段名','类型','说明'], [
    ('id', 'BIGINT (PK, IDENTITY)', '用户ID，自增主键'),
    ('username', 'VARCHAR(50)', '用户名，唯一索引'),
    ('password', 'VARCHAR(200)', 'BCrypt加密后的密码密文'),
    ('real_name', 'VARCHAR(50)', '真实姓名'),
    ('phone', 'VARCHAR(20)', '手机号'),
    ('role', 'VARCHAR(20)', '角色：OWNER/ADMIN/STAFF'),
    ('status', 'INT', '状态：1=启用 0=禁用'),
    ('avatar_url', 'VARCHAR(500)', '头像URL路径'),
    ('email', 'VARCHAR(100)', '邮箱'),
    ('gender', 'VARCHAR(10)', '性别'),
    ('birthday', 'DATE', '生日'),
    ('id_card', 'VARCHAR(20)', '身份证号'),
    ('emergency_contact', 'VARCHAR(50)', '紧急联系人'),
    ('emergency_phone', 'VARCHAR(20)', '紧急联系电话'),
    ('create_time', 'DATETIME', '创建时间'),
])

H('3.4.2 community_house（房屋表）', 3)
tbl(['字段名','类型','说明'], [
    ('id', 'BIGINT (PK, IDENTITY)', '房屋ID'),
    ('unit_id', 'BIGINT (FK)', '所属单元ID'),
    ('room_no', 'VARCHAR(20)', '房间号'),
    ('area', 'DECIMAL(10,2)', '面积（平方米）'),
    ('status', 'VARCHAR(20)', '状态：VACANT=空置 / OCCUPIED=已入住'),
    ('owner_id', 'BIGINT (FK, NULLABLE)', '绑定业主ID，NULL表示未绑定'),
    ('create_time', 'DATETIME', '创建时间'),
])

H('3.4.3 parking_space（车位表）', 3)
tbl(['字段名','类型','说明'], [
    ('id', 'BIGINT (PK, IDENTITY)', '车位ID'),
    ('space_no', 'VARCHAR(30)', '车位编号，如 A01-101'),
    ('building_id', 'BIGINT (FK)', '所属楼栋ID'),
    ('zone', 'VARCHAR(20)', '区域标识'),
    ('row_no', 'INT', '行号（3D可视化布局用）'),
    ('col_no', 'INT', '列号（3D可视化布局用）'),
    ('area', 'DECIMAL(8,2)', '面积（平方米）'),
    ('type', 'VARCHAR(20)', '类型：STANDARD/COMPACT/LARGE/VIP'),
    ('price', 'DECIMAL(10,2)', '挂牌价格'),
    ('status', 'VARCHAR(20)', '状态：FREE/SOLD/LOCKED/RESERVED'),
    ('owner_id', 'BIGINT (FK, NULLABLE)', '购买业主ID'),
    ('purchase_price', 'DECIMAL(10,2)', '实际购买价'),
    ('purchase_time', 'DATETIME', '购买时间'),
])

H('3.4.4 payment_bill（缴费账单表）', 3)
tbl(['字段名','类型','说明'], [
    ('id', 'BIGINT (PK, IDENTITY)', '账单ID'),
    ('owner_id', 'BIGINT (FK)', '业主ID'),
    ('house_id', 'BIGINT (FK, NULLABLE)', '房屋ID（房屋类费用）'),
    ('parking_space_id', 'BIGINT (FK, NULLABLE)', '车位ID（车位类费用）'),
    ('fee_item_id', 'BIGINT (FK)', '费用项目ID'),
    ('amount', 'DECIMAL(10,2)', '应缴金额'),
    ('status', 'VARCHAR(20)', '状态：PENDING/PAID/OVERDUE'),
    ('due_date', 'DATE', '缴费截止日期'),
    ('pay_time', 'DATETIME', '实际支付时间'),
    ('create_time', 'DATETIME', '生成时间'),
])

H('3.4.5 access_card（门禁卡表）', 3)
tbl(['字段名','类型','说明'], [
    ('id', 'BIGINT (PK, IDENTITY)', '门禁卡ID'),
    ('card_no', 'VARCHAR(30)', '卡片编号（AC+日期+随机数）'),
    ('owner_id', 'BIGINT (FK)', '持卡业主ID'),
    ('card_type', 'VARCHAR(20)', '类型：OWNER/FAMILY/VISITOR/TEMPORARY'),
    ('status', 'VARCHAR(20)', '状态：ACTIVE/SUSPENDED/CANCELLED'),
    ('building_ids', 'VARCHAR(200)', '可通行楼栋ID列表，逗号分隔'),
    ('valid_from', 'DATE', '有效期开始'),
    ('valid_to', 'DATE (NULLABLE)', '有效期结束，NULL=永久'),
    ('remark', 'VARCHAR(200)', '备注'),
])

H('3.4.6 其他核心表', 3)
tbl(['表名','核心字段','说明'], [
    ('facility', 'name, category, image_url, status, deposit', '社区公共设施'),
    ('facility_booking', 'facility_id, owner_id, purpose, status, handler_id', '借用申请及审批'),
    ('repair_request', 'owner_id, category, title, description, image_url, status', '设施报修工单'),
    ('feedback', 'owner_id, type, title, content, status, handler_id', '留言反馈'),
    ('feedback_reply', 'feedback_id, user_id, user_role, content', '反馈回复时间线'),
    ('access_log', 'card_no, user_name, direction, gate_location, building_id, access_status', '门禁进出记录'),
    ('announcement', 'title, content, type, publisher_id, status', '社区公告'),
    ('property_fee_item', 'item_name, amount, cycle, status', '物业费用项目'),
    ('community_building', 'building_no, name, total_floors', '楼栋信息'),
    ('community_unit', 'building_id, unit_no', '单元信息'),
])

doc.add_page_break()

# ══════════════════════════════════════════════════
#  第四章 系统设计与实现
# ══════════════════════════════════════════════════
H('四、系统设计与实现', 1)

H('4.1 系统设计目标', 2)
body('本系统的设计目标如下：')
numbered([
    '高可用性：系统采用前后端分离架构，支持独立部署和水平扩展；',
    '安全性：基于 JWT 无状态认证 + BCrypt 密码加密，实现用户身份验证和角色级权限控制；',
    '可视化：引入 Three.js 3D 渲染引擎和 ECharts 数据可视化库，提升数据展示效果；',
    '规范化：统一 RESTful API 设计规范，统一业务编号规则，统一状态机流转模型；',
    '可维护性：采用分层架构（Controller → Service → Mapper），业务逻辑清晰，易于维护。',
])

H('4.2 系统技术选型', 2)

H('4.2.1 程序架构', 3)
body('本系统采用 B/S（Browser/Server）架构，前后端分离设计：')
bullet([
    '前端：独立 SPA 应用，通过 Vite 开发服务器运行，生产环境构建为静态资源部署',
    '后端：Spring Boot RESTful API 服务，提供 JSON 接口',
    '通信：前端通过 Axios HTTP 客户端向后端发送请求，响应格式统一为 { code, msg, data }',
    '认证：JWT Token 通过 HTTP Header (Authorization: Bearer <token>) 传递',
])

H('4.2.2 技术栈详细选型', 3)

tbl(['层级','技术/框架','版本','说明'], [
    ('操作系统', 'Windows 11', '—', '开发与部署环境'),
    ('后端语言', 'Java', '17', 'LTS 版本，支持 Records 等新特性'),
    ('后端框架', 'Spring Boot', '3.1.5', 'RESTful API 服务框架'),
    ('ORM 框架', 'MyBatis-Plus', '3.5.8', '增强 MyBatis，支持 CRUD 自动生成'),
    ('数据库', 'SQL Server', '—', 'Microsoft 关系型数据库'),
    ('安全框架', 'JWT (jjwt)', '0.12.3', 'JSON Web Token 无状态认证'),
    ('密码加密', 'Spring Security Crypto', '—', 'BCrypt 单向加密'),
    ('工具库', 'Lombok', '1.18.38', '编译期代码生成，简化 POJO'),
    ('前端框架', 'Vue.js', '3.5.32', '渐进式 JavaScript 框架（Composition API）'),
    ('构建工具', 'Vite', '8.0.4', '下一代前端构建工具'),
    ('UI 组件库', 'Element Plus', '2.13.7', '基于 Vue 3 的桌面端 UI 组件库'),
    ('HTTP 客户端', 'Axios', '1.15.0', '基于 Promise 的 HTTP 请求库'),
    ('图表库', 'ECharts', '6.0.0', 'Apache 开源数据可视化图表库'),
    ('3D 引擎', 'Three.js', '0.184.0', 'WebGL 3D 渲染引擎'),
    ('路由', 'Vue Router', '4.6.4', 'Vue.js 官方路由管理器'),
    ('富文本编辑', 'WangEditor', '5.1.23', '轻量级富文本编辑器'),
    ('预处理器', 'Sass', '1.99.0', 'CSS 预处理器'),
])

H('4.3 系统功能模块设计', 2)

H('4.3.1 整体模块划分', 3)
body('系统按角色划分为两大功能域，共计 22 个功能页面/接口组：')

code('''
智慧社区物业管理系统
├── 公共模块
│   ├── 用户注册（Register）
│   ├── 用户登录（Login）
│   └── JWT 认证（JwtAuthFilter）
│
├── 业主端（Owner）
│   ├── 首页概览（OwnerHome）
│   ├── 在线缴费（OwnerBills）
│   ├── 车位选购（OwnerParking）— 3D可视化
│   ├── 设施借用（OwnerFacility）
│   ├── 设施报修（OwnerRepair）
│   ├── 留言反馈（OwnerFeedback）
│   ├── 门禁卡查看（OwnerAccessCard）
│   ├── 社区公告（OwnerAnnouncement）
│   └── 个人中心（OwnerProfile）
│
└── 管理端（Admin）
    ├── 数据概览（Dashboard）— ECharts
    ├── 用户管理（UserManage）
    ├── 费用项目管理（FeeManage）
    ├── 账单管理（AdminBill）
    ├── 楼栋房屋管理（BuildingHouseManage）
    ├── 楼栋平面图（BuildingVisual）
    ├── 车位管理（AdminParkingManage）
    ├── 设施管理（AdminFacilityManage）
    ├── 报修管理（RepairManage）
    ├── 留言反馈管理（AdminFeedbackManage）
    ├── 门禁卡管理（AdminAccessCardManage）
    ├── 进出记录（AdminAccessLog）
    └── 公告管理（AnnouncementManage）
''')

H('4.3.2 后端接口设计', 3)
body('系统后端共设计 17 个 Controller 类，提供约 80+ 个 RESTful API 接口。以下为核心接口分组：')

tbl(['模块','接口前缀','核心接口','方法数'], [
    ('认证', '/api/auth', 'login, register, info', '3'),
    ('用户管理', '/api/users', 'page, add, update, delete, status, bind-house', '10+'),
    ('缴费账单', '/api/owner/bills', 'list, pay', '2'),
    ('账单管理', '/api/admin/bills', 'page, generate, eligible, status', '6'),
    ('费用项目', '/api/admin/fee-items', 'page, add, update, delete', '4'),
    ('车位', '/api/parking', 'buildings, visual, grid, purchase, my, admin CRUD', '15+'),
    ('设施', '/api/facility', 'page, detail, categories, stats, CRUD', '8'),
    ('设施借用', '/api/facility-booking', 'submit, mine, return, page, handle', '5'),
    ('报修', '/api/repair', 'page, submit, handle', '3'),
    ('反馈', '/api/feedback', 'page, submit, reply, close, replies', '5'),
    ('门禁卡', '/api/access-card', 'page, my, issue, edit, suspend, resume, cancel, delete, stats', '9'),
    ('进出记录', '/api/access-log', 'page, stats', '2'),
    ('楼栋', '/api/buildings', 'page, add, update, delete', '4'),
    ('单元', '/api/units', 'page, by-building, add, update, delete', '5'),
    ('房屋', '/api/houses', 'page, checkin, checkout, map-data, CRUD', '8'),
    ('公告', '/api/announcements', 'page, detail, CRUD, publish, unpublish', '8'),
    ('文件上传', '/api/upload', 'image, facility', '2'),
])

H('4.3.3 前端页面设计', 3)
body('前端采用 Vue 3 Composition API 开发，共 26 个 .vue 组件文件，按功能域组织为以下结构：')

tbl(['区域','页面','文件','行数'], [
    ('公共', '登录页', 'Login.vue', '433'),
    ('公共', '注册页', 'Register.vue', '266'),
    ('业主', '布局框架', 'owner/Layout.vue', '410'),
    ('业主', '首页', 'OwnerHome.vue', '774'),
    ('业主', '在线缴费', 'OwnerBills.vue', '304'),
    ('业主', '车位选购', 'OwnerParking.vue', '412'),
    ('业主', '设施借用', 'OwnerFacility.vue', '401'),
    ('业主', '设施报修', 'OwnerRepair.vue', '231'),
    ('业主', '留言反馈', 'OwnerFeedback.vue', '443'),
    ('业主', '个人中心', 'OwnerProfile.vue', '675'),
    ('管理', '布局框架', 'admin/Layout.vue', '436'),
    ('管理', '数据概览', 'Dashboard.vue', '346'),
    ('管理', '用户管理', 'UserManage.vue', '310'),
    ('管理', '费用管理', 'FeeManage.vue', '345'),
    ('管理', '楼栋房屋', 'BuildingHouseManage.vue', '367'),
    ('管理', '车位管理', 'AdminParkingManage.vue', '598'),
    ('管理', '设施管理', 'AdminFacilityManage.vue', '336'),
])

H('4.4 系统安全设计', 2)

H('4.4.1 JWT 认证机制', 3)
body('系统采用 JWT（JSON Web Token）实现无状态认证，核心流程如下：')
numbered([
    '登录验证：用户提交用户名+密码，后端通过 BCrypt 密码比对验证身份',
    'Token 生成：验证通过后，JwtUtil 生成包含 userId（subject）和 role（claim）的 JWT Token，有效期 24 小时，签名算法为 HMAC-SHA',
    'Token 传递：前端 Axios 拦截器自动在每个请求的 Authorization Header 中注入 Bearer Token',
    'Token 验证：后端 JwtAuthFilter（OncePerRequestFilter）拦截所有请求，自动解析 Token 并将 userId/role 注入 Request 上下文',
    '路由守卫：前端 Vue Router beforeEach 钩子检查 sessionStorage 中是否存在 token，不存在则跳转登录页',
    '角色校验：登录时额外校验 role 参数与用户记录的角色是否匹配，防止越权登录',
])

H('4.4.2 密码安全', 3)
body('系统使用 Spring Security Crypto 提供的 BCryptPasswordEncoder 进行密码加密：')
bullet([
    '注册时：用户输入的明文密码经 BCrypt 加密后存储到数据库，密文包含算法标识和随机盐值',
    '登录时：使用 matches() 方法比对明文密码与数据库密文，无需解密',
    '修改密码时：先验证旧密码正确性，再对新密码进行 BCrypt 加密并更新',
    'BCrypt 特性：相同明文每次加密生成不同密文，有效防止彩虹表攻击',
])

H('4.5 系统亮点功能', 2)

H('4.5.1 车位 3D 可视化选购', 3)
body('本系统创新性地引入 Three.js WebGL 渲染引擎，实现车位 3D 可视化选购功能：')
bullet([
    '前端按楼栋加载车位网格数据（row_no × col_no），渲染为 3D 方块矩阵',
    '按车位状态着色：FREE（绿色#52c41a）、SOLD（灰色#bfbfbf）、LOCKED（红色#ff4d4f）、RESERVED（黄色#faad14）',
    '支持鼠标旋转、缩放、点击交互，点击空闲车位弹出详情和购买确认',
    '楼栋间切换自动重新渲染，实现"一幢一楼栋一车位区"的可视化架构',
])

H('4.5.2 ECharts 数据概览', 3)
body('Dashboard 页面通过 ECharts 6.0 渲染两种统计图表：')
bullet([
    '缴费状态柱状图：横轴为月份，纵轴为账单数量，三根柱形分别表示已缴（绿色）、待缴（黄色）、逾期（红色）',
    '房屋入住率饼图：环形图展示已入住（蓝色）和空置（灰色）的比例，中心显示入住率百分比',
    '图表组件采用响应式设计，在 onMounted 时初始化、onUnmounted 时销毁实例，防止内存泄漏',
])

H('4.5.3 楼栋三维可视化', 3)
body('楼栋平面图模块使用 Three.js 构建楼栋三维模型，按楼层和单元渲染房屋方块，支持入住状态着色（已入住=蓝色、空置=灰色）。管理员可直观查看各楼栋的房屋分布和入住情况。')

doc.add_page_break()

# ══════════════════════════════════════════════════
#  第五章 系统分工与合作感想
# ══════════════════════════════════════════════════
H('五、系统分工与合作感想', 1)

H('5.1 团队分工', 2)
body('本项目由三名成员协作完成，具体分工如下：')

tbl(['姓名','学号','角色','负责模块'], [
    ('蒙焕好', '2023112596', '组长',
     '用户认证（注册/登录/JWT）、管理员认证与权限验证、用户管理（CRUD+房屋绑定）、留言反馈模块（业主提交+管理端回复+状态流转）、个人中心（信息编辑/密码修改/头像上传）、社区公告模块、项目整体架构设计与协调'),
    ('农麒民', '2023112563', '组员',
     '在线缴费（业主账单查询+支付）、车位购买（3D可视化选购+购买+释放）、物业费用项目管理（CRUD）、缴费账单管理（批量生成+防重复+可收费对象查询）、车位管理（管理端CRUD+锁定/解锁+统计）、数据概览Dashboard（ECharts图表+统计卡片）'),
    ('刘向阳', '2023112568', '组员',
     '设施借用（申请+审批+归还）、设施报修（提交+工单处理）、门禁卡管理（发行+挂失+注销+全生命周期）、进出记录管理（多维查询+统计概览）、楼盘管理（楼栋→单元→房屋三级管理+入住退租+平面图可视化）、设施管理与维修记录（CRUD+软删除）'),
])

H('5.2 协作方式', 2)
body('项目采用 Git 版本控制进行代码协作，后端和前端分别作为独立仓库管理。开发过程中使用以下协作方式：')
bullet([
    '代码版本管理：Git + GitLab，各成员在独立分支开发，通过 Merge Request 合并至主分支',
    '接口约定：后端先输出 RESTful API 接口文档（路径、方法、参数、返回值），前端据此进行页面开发',
    '前后端联调：通过本地开发环境联调（后端 :8081 + 前端 Vite Dev Server），使用浏览器开发者工具排查问题',
    '定期会议：每周进行项目进度同步，讨论技术方案和遇到的问题',
])

H('5.3 合作感想', 2)

bodyb('蒙焕好（组长）：')
body('作为组长，我负责了项目的整体架构设计和技术选型。从最初确定采用 Spring Boot + Vue 3 的前后端分离架构，到引入 JWT 认证机制和 MyBatis-Plus ORM 框架，每一个技术决策都经过了充分的调研和讨论。在开发过程中，我深感团队协作的重要性——清晰的接口约定和规范的代码风格是多人协作的基础。留言反馈模块的双向沟通设计和用户管理的多表联查让我对业务建模有了更深的理解。同时，作为组长也锻炼了我的项目管理和协调能力，学会了如何在有限时间内合理分配任务、把控进度。')

bodyb('农麒民（组员）：')
body('我主要负责缴费和车位相关模块。在线缴费模块让我理解了支付系统的核心流程——账单生成、状态流转、防重复机制等。最大的挑战是车位 3D 可视化选购功能的实现，需要学习 Three.js 渲染引擎，将后端的网格数据转化为可交互的 3D 场景。通过反复调试材质、光照和相机参数，最终实现了流畅的车位选购体验。Dashboard 页面的 ECharts 图表开发也让我掌握了数据可视化的实践技巧。这次项目让我认识到，前端不仅仅是展示层，优秀的交互设计可以极大提升用户体验。')

bodyb('刘向阳（组员）：')
body('我负责的模块涵盖了设施借用、门禁管理和楼盘管理三大领域。门禁卡模块的全生命周期管理（发行→挂失→恢复→注销→删除）让我深入理解了状态机设计模式。进出记录模块的多维度查询和统计功能锻炼了我的 SQL 优化能力。楼盘管理的楼栋→单元→房屋三级层级结构设计和入住/退租的状态驱动操作，让我对复杂业务关系的建模有了更好的把握。楼栋平面图可视化功能将 Three.js 技术应用到实际业务场景，是一次很有价值的实践。通过这次项目，我深刻体会到模块化设计和代码复用的重要性。')

bodyb('总体感想：')
body('本次课程综合设计是对大学四年所学软件工程知识的一次全面实践。从需求分析、系统设计到编码实现和测试部署，我们完整经历了软件开发的全生命周期。在开发过程中，我们遇到了 SQL Server 中文编码问题、JWT Token 过期处理、Three.js 性能优化等实际挑战，但通过查阅文档、团队协作和反复调试，最终都得到了解决。这次项目不仅提升了我们的编程能力，更培养了团队协作精神、问题解决能力和工程化思维。我们相信，这些经验和能力将为未来的职业发展奠定坚实的基础。')

# ══════════════════════════════════════════════════
#  保存文档
# ══════════════════════════════════════════════════
doc.save(OUTPUT_PATH)
size_kb = os.path.getsize(OUTPUT_PATH) / 1024
print(f'课程综合报告已生成: {OUTPUT_PATH}')
print(f'文件大小: {size_kb:.1f} KB')
# -*- coding: utf-8 -*-
"""
生成课程综合报告（模板3）——智慧社区物业管理系统
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============ 全局样式设置 ============
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.5

# ============ 工具函数 ============
def set_run_font(run, font_name='宋体', size=Pt(12), bold=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, align=None, indent=None, font_name='宋体', size=Pt(12)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold=bold)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = indent
    return p

def add_body(text):
    return add_para(text, indent=Cm(0.74))

def add_body_bold(text):
    return add_para(text, indent=Cm(0.74), bold=True)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, '黑体', Pt(11), bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 灰色背景
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, '宋体', Pt(11))
    return table

def add_page_break():
    doc.add_page_break()

# ================================================================
#                          封面
# ================================================================
for _ in range(3):
    doc.add_paragraph()

add_para('应用软件综合课程设计 I', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='黑体', size=Pt(26))
doc.add_paragraph()
add_para('实验报告——课程综合报告', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='黑体', size=Pt(22))

for _ in range(3):
    doc.add_paragraph()

# 封面信息表
cover_info = [
    ('专业年级：', '软件工程专业2023级'),
    ('组    长：', '蒙焕好（2023112596）'),
    ('组    员：', '农麒民（2023112563）'),
    ('指导老师（职称）：', '韩敏（教授）'),
    ('提交日期：', '2026年6月18日'),
]
cover_table = doc.add_table(rows=len(cover_info), cols=2, style='Table Grid')
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(cover_info):
    c0 = cover_table.rows[i].cells[0]
    c0.text = ''
    r0 = c0.paragraphs[0].add_run(label)
    set_run_font(r0, '宋体', Pt(14), bold=True)
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c1 = cover_table.rows[i].cells[1]
    c1.text = ''
    r1 = c1.paragraphs[0].add_run(value)
    set_run_font(r1, '宋体', Pt(14))
    c0.width = Cm(6)
    c1.width = Cm(8)

add_page_break()

# ================================================================
#                     一、系统概述
# ================================================================
add_heading_styled('一、系统概述', level=1)

add_heading_styled('1.1 系统基本情况', level=2)
add_body('"智慧社区物业管理系统"是一个面向现代化住宅小区的综合信息化管理平台。系统采用B/S架构，'
         '基于Spring Boot 3 + Vue 3前后端分离技术栈开发，后端使用Java 17语言，数据库采用Microsoft SQL Server，'
         '前端基于Vue 3框架搭配Element Plus组件库、Three.js 3D引擎和ECharts数据可视化库。')
add_body('系统划分为"物业管理端（Admin）"和"业主端（Owner）"两大子系统，分别服务于物业管理人员和小区业主。'
         '物业管理端提供数据概览、用户管理、楼栋房屋管理、费用管理、车位管理、公告管理、报修管理、设施借用管理、'
         '留言反馈管理、门禁卡管理及进出记录追踪等功能；业主端提供首页消息聚合、在线缴费、设施报修、车位查看、'
         '社区公告、设施借用、留言反馈、门禁卡管理及个人中心等功能。')

add_heading_styled('1.2 问题定义', level=2)
add_body('随着城市化进程加快，住宅小区规模不断扩大，传统物业管理模式面临效率低下、信息不透明、沟通不畅等问题。'
         '具体表现在以下几个方面：')
add_body('（1）信息管理分散：业主信息、房屋信息、车位信息等分散在纸质档案或不同电子表格中，难以统一管理和快速检索。')
add_body('（2）缴费流程繁琐：物业费、车位管理费等费用收缴依赖线下通知和现金结算，效率低且容易产生差错。')
add_body('（3）报修响应慢：业主报修需求无法及时传达给物业，处理进度不透明，缺乏闭环跟踪机制。')
add_body('（4）设施管理困难：社区公共设施借用缺乏信息化手段，难以追踪借用状态和归还情况。')
add_body('（5）安全管控薄弱：传统门禁缺乏数字化管理，无法有效追踪人员进出记录和门禁卡生命周期。')
add_body('（6）信息可视化不足：楼栋、房屋、车位等空间信息缺乏直观的可视化展示手段，不便于管理和决策。')
add_body('本系统旨在通过信息化手段解决上述问题，构建一个功能完善、操作便捷、安全可靠的智慧社区综合管理平台。')

add_page_break()

# ================================================================
#                    二、系统需求分析
# ================================================================
add_heading_styled('二、系统需求分析', level=1)

add_heading_styled('2.1 系统整体业务流程', level=2)
add_body('系统的核心业务流程围绕"物业管理"和"业主服务"两条主线展开。物业管理员通过后台进行基础数据维护（楼栋、房屋、'
         '车位、用户等），生成费用账单并跟踪处理状态；业主通过前端门户查看个人信息、缴纳费用、提交报修和反馈、'
         '浏览公告、借用设施及管理门禁卡。两端通过RESTful API实时交互，数据共享于同一数据库，确保信息一致性。')

add_heading_styled('2.2 功能模块分析', level=2)

add_body_bold('（1）用户认证与权限管理模块')
add_body('系统支持用户注册、登录和角色区分（ADMIN管理员/OWNER业主）。采用JWT令牌机制进行身份认证，'
         '前端路由守卫根据Token拦截未授权访问。管理员可管理系统内所有用户信息，包括增删改查、状态启停、角色分配等。')

add_body_bold('（2）楼栋与房屋管理模块')
add_body('系统以"楼栋 → 单元 → 房屋"三级层次结构组织空间数据。管理员可管理楼栋信息、单元信息和房屋信息，'
         '支持按房间号模糊搜索。房屋具有VACANT（空置）和OCCUPIED（已入住）两种状态，管理员可执行入住和退房操作，'
         '系统将房屋与业主自动关联。前端提供基于Three.js的3D楼栋可视化模型和平面视图两种展示方式。')

add_body_bold('（3）在线缴费模块')
add_body('管理员可配置费用项目（如物业管理费、垃圾处理费等），设置收费周期（月/季/年），'
         '并按房屋维度批量生成账单，同时支持车位管理费按车位维度生成。账单编号采用"日期-业主编号"格式自动生成。'
         '业主端可查看待缴/已缴/逾期账单，并进行在线缴费操作。账单状态包含PENDING（待缴）、PAID（已缴）、OVERDUE（逾期）三种。')

add_body_bold('（4）车位管理模块')
add_body('车位具有STANDARD/COMPACT/LARGE/VIP四种类型和FREE/LOCKED/SOLD/RESERVED四种状态。'
         '管理员可管理车位信息、设置价格、执行出售/预订/锁定操作。车位与楼栋关联，前端提供3D可视化地图展示车位分布，'
         '已售车位显示业主信息。业主端可查看自己名下的车位详情。')

add_body_bold('（5）报修管理模块')
add_body('业主可提交报修申请，选择报修类别、填写标题和描述，支持上传图片。报修单号自动生成，格式为"日期-R+业主编号"。'
         '管理员可查看报修列表、指派处理人、回复处理结果。报修状态包含PENDING（待处理）、PROCESSING（处理中）、'
         'COMPLETED（已完成）等状态，实现报修全流程闭环管理。')

add_body_bold('（6）社区公告模块')
add_body('管理员可发布、编辑、撤回和删除社区公告，支持公告分类（NOTICE通知/ACTIVITY活动/MAINTENANCE维护/OTHER其他）'
         '和置顶功能。公告状态包含DRAFT（草稿）、PUBLISHED（已发布）、WITHDRAWN（已撤回）。'
         '业主端可查看已发布的公告列表，支持按类型筛选。')

add_body_bold('（7）设施借用模块')
add_body('管理员管理社区公共设施信息（运动器材、工具设备、文娱用品等），设置借用押金。'
         '业主可浏览设施列表（含图片、状态、位置），提交借用申请。管理员审批通过后，业主可在约定时间借用并归还。'
         '设施状态包含AVAILABLE（可借用）、BOOKED（已借出）、MAINTENANCE（维护中）、RETIRED（已报废）。'
         '借用申请状态包含PENDING/APPROVED/REJECTED/RETURNED四种。')

add_body_bold('（8）留言反馈模块')
add_body('业主可提交留言反馈（SUGGESTION建议/COMPLAINT投诉/INQUIRY咨询），支持上传多张图片。'
         '反馈单号自动生成，格式为"日期-F+业主编号+记录编号"。管理员和业主可进行多轮对话式回复。'
         '状态包含PENDING（待处理）、PROCESSING（处理中）、REPLIED（已回复）、CLOSED（已关闭）。')

add_body_bold('（9）门禁卡管理模块')
add_body('系统支持四种门禁卡类型：OWNER业主卡、FAMILY家庭卡、VISITOR访客卡、TEMPORARY临时卡。'
         '门禁卡具有完整生命周期：发行 → 编辑权限 → 挂失 → 恢复/注销 → 删除。'
         '每张卡可配置通行楼栋范围（逗号分隔楼栋ID，空表示全部通行）和有效期。'
         '管理员可通过管理界面执行全流程操作，业主端可查看自己的门禁卡并一键挂失。')

add_body_bold('（10）进出记录追踪模块')
add_body('系统记录每次门禁刷卡事件，包括卡号、用户、进出方向（IN进入/OUT离开）、门禁位置、'
         '通行状态（SUCCESS放行/DENIED拒绝）等信息。管理员可按时间、卡号、楼栋等条件查询进出记录，'
         '实现社区人员流动的数字化追踪。')

add_body_bold('（11）个人中心模块')
add_body('业主可查看和编辑个人信息，包括姓名、手机号、邮箱、性别、生日、身份证号、紧急联系人等字段，'
         '支持头像上传。同时展示业主的房屋关联信息（楼栋、单元、房间号）。')

add_body_bold('（12）数据概览模块（Dashboard）')
add_body('管理端首页提供数据总览面板，基于ECharts图表库展示关键运营指标：用户统计、房屋入住率、'
         '账单收缴率、报修处理情况等，以柱状图、饼图、进度条等形式直观呈现，辅助管理决策。')

add_page_break()

# ================================================================
#                   三、数据库设计
# ================================================================
add_heading_styled('三、数据库设计', level=1)

add_heading_styled('3.1 数据流图分析', level=2)
add_body('系统的数据流围绕"用户"这一核心角色展开。物业管理员作为数据的维护者和审批者，'
         '向系统输入楼栋、房屋、车位、费用项目等基础数据，并处理业主提交的各类申请；'
         '业主作为数据的使用者和申请者，从系统获取账单、公告、设施等信息，并提交缴费、报修、反馈、借用等业务请求。'
         '系统内部通过JWT令牌进行身份验证和权限校验，确保数据在不同角色间安全流转。')
add_body('核心数据流如下：')
add_body('① 管理员 → 基础数据（楼栋/房屋/车位/费用项） → 数据库')
add_body('② 管理员 → 账单生成 → 数据库 → 业主端展示')
add_body('③ 业主 → 报修/反馈/借用申请 → 数据库 → 管理员审批')
add_body('④ 管理员 → 审批回复 → 数据库 → 业主端展示')
add_body('⑤ 业主 → 缴费操作 → 数据库更新账单状态')
add_body('⑥ 门禁卡操作 → 数据库 → 进出记录生成')

add_heading_styled('3.2 实体关系模型', level=2)
add_body('系统共包含15个核心实体，各实体及其主要属性如下表所示：')

# 实体表
add_table(
    ['实体名称', '对应数据表', '主要属性'],
    [
        ['系统用户', 'sys_user', 'id, username, password, realName, phone, role, email, gender, birthday, idCard, avatarUrl, status'],
        ['楼栋', 'community_building', 'id, buildingNo, name, totalFloors'],
        ['单元', 'community_unit', 'id, buildingId, unitNo'],
        ['房屋', 'community_house', 'id, unitId, roomNo, area, status(VACANT/OCCUPIED), ownerId'],
        ['车位', 'parking_space', 'id, spaceNo, buildingId, zone, rowNo, colNo, area, type, status, ownerId, price'],
        ['费用项目', 'property_fee_item', 'id, itemName, amount, cycle(MONTH/QUARTER/YEAR), status'],
        ['账单', 'payment_bill', 'id, ownerId, houseId, parkingSpaceId, feeItemId, amount, status, dueDate, payTime'],
        ['报修申请', 'repair_request', 'id, ownerId, category, title, description, imageUrl, status, handlerId, replyContent'],
        ['公告', 'announcement', 'id, title, content, type, publisherId, status, isTop'],
        ['设施', 'facility', 'id, name, category, description, imageUrl, location, deposit, status'],
        ['设施借用', 'facility_booking', 'id, facilityId, ownerId, purpose, durationHours, status, handlerId, returnTime'],
        ['留言反馈', 'feedback', 'id, ownerId, type, title, content, images, status, handlerId'],
        ['反馈回复', 'feedback_reply', 'id, feedbackId, userId, userRole, content'],
        ['门禁卡', 'access_card', 'id, cardNo, ownerId, cardType, status, buildingIds, validFrom, validTo'],
        ['进出记录', 'access_log', 'id, cardId, cardNo, userId, userName, direction, gateLocation, buildingId, accessTime, accessStatus'],
    ]
)

doc.add_paragraph()
add_body('实体间的主要关系如下：')
add_body('（1）楼栋与单元：一对多关系。一栋楼包含多个单元，单元通过buildingId外键关联楼栋。')
add_body('（2）单元与房屋：一对多关系。一个单元包含多间房屋，房屋通过unitId外键关联单元。')
add_body('（3）用户与房屋：多对多关系（通过ownerId实现）。一个业主可拥有多间房屋，一间房屋对应一个业主。')
add_body('（4）用户与车位：一对多关系。一个业主可拥有多个车位，车位通过ownerId关联业主。')
add_body('（5）车位与楼栋：多对一关系。车位通过buildingId关联所属楼栋。')
add_body('（6）用户与账单：一对多关系。账单通过ownerId关联业主，同时关联房屋（houseId）和费用项目（feeItemId）。')
add_body('（7）用户与报修申请：一对多关系。报修通过ownerId关联提交人，通过handlerId关联处理人。')
add_body('（8）设施与借用申请：一对多关系。借用申请通过facilityId关联设施，通过ownerId关联申请人。')
add_body('（9）留言反馈与回复：一对多关系。回复通过feedbackId关联反馈记录。')
add_body('（10）用户与门禁卡：一对多关系。门禁卡通过ownerId关联合计业主。')
add_body('（11）门禁卡与进出记录：一对多关系。进出记录通过cardId关联门禁卡。')

add_page_break()

# ================================================================
#                  四、系统设计与实现
# ================================================================
add_heading_styled('四、系统设计与实现', level=1)

add_heading_styled('4.1 系统设计目标', level=2)
add_body('本系统的设计目标如下：')
add_body('（1）功能完备性：覆盖物业管理的核心业务场景，包括用户管理、房屋管理、车位管理、费用收缴、'
         '报修处理、公告发布、设施借用、留言反馈、门禁管控等，形成完整的物业管理信息化解决方案。')
add_body('（2）双端协同：系统同时提供物业管理端和业主端，两端共享同一数据源，通过RESTful API实时交互，'
         '实现管理侧与服务侧的无缝协同。')
add_body('（3）可视化直观：采用Three.js 3D引擎实现楼栋和车位的三维可视化展示，采用ECharts实现运营数据的图表化展示，'
         '提升用户体验和管理效率。')
add_body('（4）安全可靠：采用JWT令牌认证机制，前端路由守卫拦截未授权访问，后端接口层进行角色权限校验，'
         '保障系统数据安全。')
add_body('（5）可扩展性：采用模块化分层架构（Controller → Service → Mapper），各模块低耦合高内聚，'
         '便于后续功能扩展和维护。')

add_heading_styled('4.2 系统选型', level=2)

add_table(
    ['选型维度', '技术方案', '说明'],
    [
        ['程序架构', 'B/S架构（前后端分离）', '后端提供RESTful API，前端通过HTTP请求交互'],
        ['后端开发语言', 'Java 17', 'LTS版本，支持Records、Pattern Matching等现代语法'],
        ['后端框架', 'Spring Boot 3.1.5', '主流微服务框架，自动配置、内嵌Tomcat'],
        ['ORM框架', 'MyBatis-Plus 3.5.8', '增强版MyBatis，支持代码生成、CRUD封装、自定义SQL'],
        ['数据库', 'Microsoft SQL Server', '企业级关系数据库，支持高并发和事务管理'],
        ['安全认证', 'JWT (jjwt 0.12.3)', '无状态令牌认证，支持跨域访问'],
        ['前端框架', 'Vue 3.5 + Vite 8', 'Composition API + 高性能构建工具'],
        ['UI组件库', 'Element Plus 2.13', 'Vue 3生态主流UI库，组件丰富、文档完善'],
        ['3D可视化', 'Three.js 0.184', 'WebGL 3D引擎，实现楼栋和车位三维展示'],
        ['数据可视化', 'ECharts 6.0', '百度开源图表库，支持多种图表类型'],
        ['HTTP客户端', 'Axios 1.15', 'Promise风格的HTTP请求库，支持拦截器'],
        ['富文本编辑器', 'WangEditor 5.1', '轻量级富文本编辑器，用于公告编辑'],
        ['辅助工具', 'Lombok 1.18.38', '编译期代码生成，减少样板代码'],
        ['操作系统', 'Windows 11', '开发与部署环境'],
    ]
)

add_heading_styled('4.3 系统功能设计', level=2)

add_body('系统采用前后端分离的B/S架构，后端运行于8081端口，前端开发服务器通过Vite代理转发API请求。'
         '前端使用Vue Router实现路由管理，通过路由守卫和sessionStorage中的JWT令牌实现访问控制。')

doc.add_paragraph()
add_body_bold('4.3.1 物业管理端功能模块')
add_body('物业管理端面向物业工作人员，提供完整的后台管理功能，主要包括以下模块：')

add_table(
    ['功能模块', '主要功能', '对应前端页面'],
    [
        ['数据概览', '用户统计、房屋入住率、账单收缴率、报修数据等ECharts图表展示', 'Dashboard.vue'],
        ['用户管理', '用户列表、新增/编辑/删除用户、角色分配、状态启停', 'UserManage.vue'],
        ['楼栋房屋管理', '楼栋/单元/房屋三级管理，入住/退房操作，房间号搜索', 'BuildingHouseManage.vue'],
        ['楼栋平面图', '楼栋3D可视化模型（Three.js），房屋状态颜色编码，楼层交互', 'BuildingVisual.vue'],
        ['费用管理', '费用项目配置（物业费、垃圾费等），按房屋/车位维度生成账单', 'FeeManage.vue'],
        ['车位管理', '车位CRUD，出售/预订/锁定操作，3D车位地图', 'AdminParkingManage.vue'],
        ['公告管理', '公告发布/编辑/撤回/删除，分类管理，置顶功能', 'AnnouncementManage.vue'],
        ['报修管理', '报修列表查看、指派处理人、回复处理结果', 'RepairManage.vue'],
        ['设施借用管理', '设施CRUD、审批借用申请、归还确认', 'AdminFacilityManage.vue'],
        ['留言反馈管理', '反馈列表、多轮回复、状态流转', 'AdminFeedbackManage.vue'],
        ['门禁卡管理', '卡片发行/编辑/挂失/恢复/注销/删除，权限配置', 'AdminAccessCardManage.vue'],
        ['进出记录追踪', '进出记录查询，按卡号/楼栋/时间筛选', 'AdminAccessLog.vue'],
    ]
)

doc.add_paragraph()
add_body_bold('4.3.2 业主端功能模块')
add_body('业主端面向小区业主，提供便捷的生活服务功能：')

add_table(
    ['功能模块', '主要功能', '对应前端页面'],
    [
        ['业主首页', '消息提醒聚合（账单/报修/公告/反馈/预约），个人信息卡片', 'OwnerHome.vue'],
        ['在线缴费', '账单列表查看、按状态筛选、在线缴费操作', 'OwnerBills.vue'],
        ['设施报修', '提交报修申请、上传图片、查看处理进度', 'OwnerRepair.vue'],
        ['车位管理', '查看个人车位信息、3D车位地图浏览', 'OwnerParking.vue'],
        ['社区公告', '公告列表浏览、按类型筛选、公告详情查看', 'OwnerAnnouncement.vue'],
        ['设施借用', '设施浏览（图片+状态）、提交借用申请、查看借用记录', 'OwnerFacility.vue'],
        ['留言反馈', '提交反馈（建议/投诉/咨询）、上传图片、多轮对话', 'OwnerFeedback.vue'],
        ['门禁卡', '查看个人门禁卡（仿真卡面）、一键挂失', 'OwnerAccessCard.vue'],
        ['个人中心', '个人信息查看与编辑、头像上传、房屋关联信息展示', 'OwnerProfile.vue'],
    ]
)

add_heading_styled('4.3.3 后端API架构', level=3)
add_body('后端采用Controller → Service → Mapper三层架构，所有API以/api为统一前缀。'
         '系统共包含17个Controller类，提供的主要API接口如下：')

add_table(
    ['Controller', '职责', '主要接口'],
    [
        ['AuthController', '用户认证', 'POST /auth/login, POST /auth/register, GET /auth/info'],
        ['UserController', '用户管理', 'GET/POST/PUT/DELETE /users/*, GET /users/owner/profile'],
        ['CommunityBuildingController', '楼栋管理', 'GET/POST/PUT/DELETE /buildings/*'],
        ['CommunityUnitController', '单元管理', 'GET/POST/PUT/DELETE /units/*'],
        ['CommunityHouseController', '房屋管理', 'GET/POST/PUT /houses/*, POST /houses/{id}/occupy, POST /houses/{id}/vacate'],
        ['ParkingSpaceController', '车位管理', 'GET/POST/PUT/DELETE /parking/*, GET /parking/visual/grid'],
        ['AdminBillController', '账单管理', 'POST /bills/generate, GET /bills/*, PUT /bills/{id}/pay'],
        ['PaymentBillController', '业主账单', 'GET /owner/bills, PUT /owner/bills/{id}/pay'],
        ['FeeItemController', '费用项目', 'GET/POST/PUT/DELETE /fee-items/*'],
        ['RepairRequestController', '报修管理', 'GET/POST/PUT /repairs/*'],
        ['AnnouncementController', '公告管理', 'GET/POST/PUT/DELETE /announcements/*'],
        ['FacilityController', '设施管理', 'GET/POST/PUT/DELETE /facilities/*'],
        ['FacilityBookingController', '设施借用', 'GET/POST/PUT /facility-bookings/*'],
        ['FeedbackController', '留言反馈', 'GET/POST/PUT /feedbacks/*, GET/POST /feedbacks/{id}/replies'],
        ['AccessCardController', '门禁卡管理', 'GET/POST/PUT/DELETE /access-cards/*'],
        ['AccessLogController', '进出记录', 'GET /access-logs, GET /access-logs/stats'],
        ['FileUploadController', '文件上传', 'POST /upload/repair, POST /upload/facility, POST /upload/feedback, POST /upload/avatar'],
    ]
)

add_heading_styled('4.3.4 前端路由与页面架构', level=3)
add_body('前端采用Vue Router 4实现路由管理，路由结构如下：')
add_body('• /login — 登录页（建筑蓝图风格SVG背景）')
add_body('• /register — 注册页')
add_body('• /owner — 业主端布局（左侧导航栏 + 内容区），包含9个子页面')
add_body('• /admin — 管理端布局（顶部导航栏 + 左侧菜单 + 内容区），包含12个子页面')
add_body('前端通过Axios封装HTTP请求，统一设置Bearer Token认证头和错误拦截器。'
         'Vite开发服务器配置反向代理，将/api请求转发至后端8081端口，解决开发环境跨域问题。')

add_heading_styled('4.3.5 关键技术实现', level=3)
add_body('（1）3D楼栋可视化：基于Three.js WebGL引擎，将每间房渲染为独立3D方块，按楼层×单元立体排布。'
         '支持鼠标拖拽旋转、滚轮缩放、点击查看详情。颜色编码区分房屋状态（绿色=已入住、橙色=已出租、灰色=空置）。')
add_body('（2）3D车位地图：采用Three.js实现车位三维可视化，支持车位选择、状态展示、业主信息悬浮提示。')
add_body('（3）消息聚合引擎：业主首页聚合5类系统消息（账单、报修、公告、反馈、设施预约），'
         '支持红点计数、类型图标、时间展示和点击跳转。')
add_body('（4）业务单据编号自动生成：账单（日期-P+业主编号）、报修（日期-R+业主编号）、反馈（日期-F+业主编号+记录编号），'
         '均通过实体类计算方法实现，无需数据库额外字段。')
add_body('（5）JWT认证与权限控制：登录成功生成JWT令牌，前端存储于sessionStorage，'
         '每次请求通过Axios拦截器自动附加Authorization头。后端通过自定义拦截器解析令牌并注入userId。')

add_page_break()

# ================================================================
#                  五、系统分工与合作感想
# ================================================================
add_heading_styled('五、系统分工与合作感想', level=1)

add_heading_styled('5.1 系统分工', level=2)

add_table(
    ['成员', '学号', '角色', '负责模块'],
    [
        ['蒙焕好', '2023112596', '组长',
         '用户认证与权限管理、用户管理、留言反馈（前后端）、门禁卡管理（前后端）、进出记录追踪、个人中心、登录注册页面、系统架构设计与技术选型、数据库设计与数据初始化'],
        ['农麒民', '2023112563', '组员',
         '在线缴费模块（前后端）、车位管理模块（前后端）、3D车位可视化、数据概览Dashboard（ECharts图表）、楼栋房屋管理、楼栋3D可视化、公告管理、报修管理、设施借用管理'],
    ]
)

add_heading_styled('5.2 合作感想', level=2)
add_body('在本次课程设计中，我们小组两人协作完成了"智慧社区物业管理系统"的全栈开发，从需求分析、数据库设计、'
         '后端API开发、前端页面构建到3D可视化实现，完整经历了一个软件项目从0到1的全过程。以下是我们的几点合作体会：')

add_body('第一，明确分工与紧密协作缺一不可。我们将系统按模块划分为两大块，每人负责若干模块的前后端开发，'
         '这样保证了并行开发的效率；同时，在涉及跨模块联调（如账单与车位的数据联动、门禁卡与进出记录的关系）时，'
         '我们会及时沟通接口规范和数据结构，确保系统整体一致性。')

add_body('第二，统一的技术规范至关重要。项目初期我们约定了统一的API命名规范（RESTful风格）、'
         '数据返回格式（Result统一封装）、前端组件命名规则等，这大大减少了后期集成时的摩擦和返工。')

add_body('第三，技术挑战推动了学习成长。本项目中我们首次引入了Three.js 3D可视化技术来实现楼栋和车位的三维展示，'
         '从最初对WebGL一无所知到最终实现可交互的3D场景，这一过程极大地锻炼了我们的自学能力和问题解决能力。'
         '同时，SQL Server与MyBatis-Plus的适配、JWT认证机制的实现、ECharts数据可视化等技术实践，'
         '都让我们对全栈开发有了更深入的理解。')

add_body('第四，测试数据的准备不可忽视。为了验证系统各模块的正确性，我们通过PowerShell脚本批量生成了'
         '楼栋、房屋、用户、车位等测试数据，并编写了多个数据校验和修复脚本。这一过程让我们认识到，'
         '真实、充分的测试数据是保障软件质量的重要基础。')

add_body('总的来说，本次课程设计让我们在实践中深化了软件工程的理论知识，体验了团队协作的力量，'
         '也积累了宝贵的项目开发经验。我们认识到，一个优秀的软件系统不仅需要扎实的技术功底，'
         '更需要清晰的业务理解、合理的架构设计和高效的团队合作。这些经验将对我们未来的职业发展产生深远的影响。')

# ============ 保存文档 ============
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '课程综合报告_智慧社区物业管理系统_2023112596_蒙焕好.docx')
doc.save(output_path)
print(f'报告已生成：{output_path}')
